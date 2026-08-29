from __future__ import annotations

import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(sys.argv[1])
NAVY = "17365D"
BLUE = "D9EAF7"
GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (
        ("Heading 1", 16, 12, 6, NAVY),
        ("Heading 2", 13, 10, 5, NAVY),
        ("Heading 3", 12, 8, 4, "1F4D78"),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_slot(doc: Document, label: str, token: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"{label}: ")
    run.bold = True
    token_run = p.add_run(token)
    token_run.font.color.rgb = RGBColor.from_string("7F6000")


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
configure_styles(doc)

header = section.header.paragraphs[0]
header.text = "{{organization_name}}  |  INTERNAL CONTROLLER MEMORANDUM"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.name = "Arial"
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = RGBColor.from_string("666666")
footer = section.footer.paragraphs[0]
footer.text = "DRAFT — NOT APPROVED"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.name = "Arial"
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor.from_string("666666")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run("INTERNAL CONTROLLER MEMORANDUM")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(10)
r.font.color.rgb = RGBColor.from_string(NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r = p.add_run("{{memo_title}}")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(24)
r.font.color.rgb = RGBColor.from_string(NAVY)

meta = doc.add_table(rows=5, cols=2)
meta.style = "Table Grid"
set_table_geometry(meta, [2700, 6660])
for row, (label, value) in zip(meta.rows, [
    ("Organization", "{{organization_name}}"),
    ("Period / as-of", "{{as_of_date}}"),
    ("Prepared by", "{{preparer_role}}"),
    ("Review owner", "{{reviewer_role}}"),
    ("Status", "DRAFT — NOT APPROVED"),
]):
    row.cells[0].text = label
    row.cells[1].text = value
    set_cell_shading(row.cells[0], BLUE)
    row.cells[0].paragraphs[0].runs[0].bold = True

add_heading(doc, "Decision requested", 1)
doc.add_paragraph("{{decision_requested}}")
add_heading(doc, "Executive recommendation", 1)
callout = doc.add_table(rows=1, cols=1)
callout.style = "Table Grid"
set_table_geometry(callout, [9360])
set_cell_shading(callout.cell(0, 0), GRAY)
callout.cell(0, 0).text = "{{executive_recommendation}}"
callout.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
add_heading(doc, "Scope and information boundary", 1)
doc.add_paragraph("{{scope_and_information_boundary}}")
doc.add_paragraph("This memorandum is supporting analysis. It does not itself evidence approval, execution, or an external party’s position.")

add_page_break(doc)
add_heading(doc, "1. Background and governing evidence", 1)
doc.add_paragraph("{{background_summary}}")
add_heading(doc, "Authoritative sources", 2)
sources = doc.add_table(rows=2, cols=4)
sources.style = "Table Grid"
set_table_geometry(sources, [1900, 2300, 2200, 2960])
for idx, label in enumerate(["Source ID", "Authority class", "Locator / citation", "What it establishes"]):
    sources.cell(0, idx).text = label
    set_cell_shading(sources.cell(0, idx), BLUE)
    sources.cell(0, idx).paragraphs[0].runs[0].bold = True
for idx, value in enumerate(["{{source_id_1}}", "{{authority_1}}", "{{source_locator_1}}", "{{source_scope_1}}"]):
    sources.cell(1, idx).text = value
add_heading(doc, "Prior-period practice", 2)
doc.add_paragraph("{{prior_period_context}}")
add_heading(doc, "Known limitations", 2)
doc.add_paragraph("{{known_limitations}}")

add_page_break(doc)
add_heading(doc, "2. Issue-by-issue analysis", 1)
for number in range(1, 4):
    add_heading(doc, f"Issue {number}: {{{{issue_{number}_title}}}}", 2)
    add_slot(doc, "Question", f"{{{{issue_{number}_question}}}}")
    add_slot(doc, "Evidence", f"{{{{issue_{number}_evidence}}}}")
    add_slot(doc, "Analysis", f"{{{{issue_{number}_analysis}}}}")
    add_slot(doc, "Conclusion", f"{{{{issue_{number}_conclusion}}}}")

add_page_break(doc)
add_heading(doc, "3. Alternatives, sensitivity, and risk", 1)
alternatives = doc.add_table(rows=4, cols=4)
alternatives.style = "Table Grid"
set_table_geometry(alternatives, [1700, 2900, 2300, 2460])
for idx, label in enumerate(["Option", "Treatment", "Evidence dependency", "Risk / sensitivity"]):
    alternatives.cell(0, idx).text = label
    set_cell_shading(alternatives.cell(0, idx), BLUE)
    alternatives.cell(0, idx).paragraphs[0].runs[0].bold = True
for row_no in range(1, 4):
    values = [f"{{{{option_{row_no}}}}}", f"{{{{treatment_{row_no}}}}}", f"{{{{evidence_{row_no}}}}}", f"{{{{risk_{row_no}}}}}"]
    for idx, value in enumerate(values):
        alternatives.cell(row_no, idx).text = value
add_heading(doc, "Recommended safeguards", 2)
doc.add_paragraph("{{recommended_safeguards}}")
add_heading(doc, "Sensitivity or threshold", 2)
doc.add_paragraph("{{sensitivity_analysis}}")

add_page_break(doc)
add_heading(doc, "4. Actions, open items, and appendices", 1)
actions = doc.add_table(rows=4, cols=5)
actions.style = "Table Grid"
set_table_geometry(actions, [1200, 3000, 1800, 1500, 1860])
for idx, label in enumerate(["Item", "Action / question", "Owner role", "Due date", "Status / evidence"]):
    actions.cell(0, idx).text = label
    set_cell_shading(actions.cell(0, idx), BLUE)
    actions.cell(0, idx).paragraphs[0].runs[0].bold = True
for row_no in range(1, 4):
    values = [f"{{{{item_{row_no}}}}}", f"{{{{action_{row_no}}}}}", f"{{{{owner_{row_no}}}}}", f"{{{{due_{row_no}}}}}", f"{{{{status_{row_no}}}}}"]
    for idx, value in enumerate(values):
        actions.cell(row_no, idx).text = value
add_heading(doc, "Required approvals (record only after they occur)", 2)
doc.add_paragraph("{{approval_record_locator_or_not_yet_obtained}}")
add_heading(doc, "Appendix index", 2)
doc.add_paragraph("{{appendix_index}}")
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)

minimal_core = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"/>'''
minimal_app = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"/>'''
with zipfile.ZipFile(OUTPUT, "r") as source:
    members = [(info, source.read(info.filename)) for info in source.infolist()]
with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=OUTPUT.parent) as temporary:
    temporary_path = Path(temporary.name)
try:
    with zipfile.ZipFile(temporary_path, "w") as target:
        for info, payload in members:
            if info.filename == "docProps/core.xml":
                payload = minimal_core
            elif info.filename == "docProps/app.xml":
                payload = minimal_app
            target.writestr(info, payload)
    temporary_path.replace(OUTPUT)
finally:
    temporary_path.unlink(missing_ok=True)
